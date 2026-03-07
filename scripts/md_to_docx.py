"""
将华工科技策略报告 Markdown 转换为格式化 Word 文档
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 颜色常量 ────────────────────────────────────────────────────────────────
C_TITLE_BG   = RGBColor(0x1A, 0x37, 0x6C)   # 深蓝封面背景
C_H1_BG      = RGBColor(0x1A, 0x37, 0x6C)   # 一级标题底色
C_H2_BG      = RGBColor(0x21, 0x5F, 0x99)   # 二级标题底色
C_H3_LINE    = RGBColor(0x21, 0x5F, 0x99)   # 三级标题左边框色
C_TBL_HEADER = RGBColor(0x21, 0x5F, 0x99)   # 表格表头底色
C_TBL_ALT    = RGBColor(0xEA, 0xF1, 0xF8)   # 表格交替行底色
C_CODE_BG    = RGBColor(0xF4, 0xF6, 0xF8)   # 代码块背景
C_ACCENT     = RGBColor(0xC0, 0x39, 0x2B)   # 强调红色
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK       = RGBColor(0x1A, 0x1A, 0x2E)
C_QUOTE_LINE = RGBColor(0x21, 0x5F, 0x99)


def rgb_hex(color: RGBColor) -> str:
    return f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    tcPr.append(shd)


def add_cell_border(cell, sides=('top', 'bottom', 'left', 'right'), color='215F99', sz=6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_para_border_left(para, color='215F99', sz=24):
    """给段落加左边框（用于 blockquote / H3）"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def set_para_shading(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    pPr.append(shd)


def apply_inline(run, text):
    """去掉 **bold** 标记，粗体由调用方控制"""
    return text


def add_run_with_bold(para, text):
    """解析行内 **bold** 并分段 run 输出"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        m = re.match(r'\*\*([^*]+)\*\*', part)
        if m:
            run = para.add_run(m.group(1))
            run.bold = True
        else:
            if part:
                para.add_run(part)


def make_document():
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style.font.color.rgb = C_DARK
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


def add_cover(doc, title_text, meta_lines):
    """封面页"""
    # 大标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(p, C_TITLE_BG)
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after  = Pt(20)
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = C_WHITE
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 元信息
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_shading(p, C_TITLE_BG)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xB3, 0xCC, 0xE8)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 尾部留白
    p = doc.add_paragraph()
    set_para_shading(p, C_TITLE_BG)
    p.paragraph_format.space_before = Pt(60)

    doc.add_page_break()


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    set_para_shading(p, C_H1_BG)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = C_WHITE
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, C_H2_BG)
    run = p.add_run(f'  {text}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_WHITE
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    set_para_border_left(p, '215F99', 24)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_H2_BG
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_ACCENT
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    add_run_with_bold(p, text)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_blockquote(doc, text):
    # 去掉 > 前缀
    text = re.sub(r'^>\s*', '', text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    set_para_border_left(p, '215F99', 20)
    set_para_shading(p, RGBColor(0xEA, 0xF1, 0xF8))
    add_run_with_bold(p, text)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x24, 0x47, 0x6E)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, C_CODE_BG)
    text = '\n'.join(lines)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '215F99')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_table_rows(lines):
    """解析 Markdown 表格为二维数组"""
    rows = []
    for line in lines:
        if re.match(r'\s*\|[-:| ]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= cols:
                break
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 清理 markdown 粗体符号，保留可见文字
            display = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
            # 去掉 emoji（仅保留 ASCII 可打印 + 中文）
            display = re.sub(r'[^\x20-\x7E\u4e00-\u9fff\uff00-\uffef\u3000-\u303f%+\-~×()（）【】《》、，。：？！…\u2014\u2015\u2019\u201c\u201d★☆]', '', display)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

            is_bold = '**' in cell_text
            run = p.add_run(display)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            if i == 0:
                # 表头行
                set_cell_bg(cell, C_TBL_HEADER)
                add_cell_border(cell, color='1A376C')
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = C_WHITE
            elif i % 2 == 0:
                set_cell_bg(cell, C_TBL_ALT)
                add_cell_border(cell, color='B3CCE8', sz=4)
                run.font.size = Pt(9)
                run.font.color.rgb = C_DARK
                if is_bold:
                    run.bold = True
            else:
                set_cell_bg(cell, RGBColor(0xFF, 0xFF, 0xFF))
                add_cell_border(cell, color='B3CCE8', sz=4)
                run.font.size = Pt(9)
                run.font.color.rgb = C_DARK
                if is_bold:
                    run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_footer(doc):
    """添加页脚"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本报告基于 AKShare MCP 实时数据生成，仅供参考，不构成投资建议。股市有风险，投资须谨慎。')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


# ── 主解析逻辑 ───────────────────────────────────────────────────────────────
def convert(md_path: str, out_path: str):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    doc = make_document()

    # 封面信息提取
    cover_title = ''
    cover_meta  = []
    for line in lines[:10]:
        stripped = line.strip()
        if stripped.startswith('# '):
            cover_title = stripped[2:]
        elif stripped.startswith('> '):
            cover_meta.append(re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped[2:]))

    add_cover(doc, cover_title, cover_meta)
    add_footer(doc)

    i = 0
    in_code  = False
    in_table = False
    code_buf = []
    tbl_buf  = []

    while i < len(lines):
        raw  = lines[i].rstrip('\n')
        line = raw.strip()
        i += 1

        # ── 代码块 ──
        if line.startswith('```'):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                add_code_block(doc, code_buf)
                in_code  = False
                code_buf = []
            continue

        if in_code:
            code_buf.append(raw)
            continue

        # ── 表格 ──
        if line.startswith('|'):
            tbl_buf.append(line)
            in_table = True
            continue
        else:
            if in_table:
                rows = parse_table_rows(tbl_buf)
                add_table(doc, rows)
                tbl_buf  = []
                in_table = False

        # ── 空行 ──
        if not line:
            continue

        # ── 分割线 ──
        if re.match(r'^---+$', line):
            add_hr(doc)
            continue

        # ── 标题 ──
        if line.startswith('#### '):
            add_h4(doc, line[5:])
        elif line.startswith('### '):
            add_h3(doc, line[4:])
        elif line.startswith('## '):
            add_h2(doc, line[3:])
        elif line.startswith('# '):
            add_h1(doc, line[2:])

        # ── blockquote ──
        elif line.startswith('> '):
            add_blockquote(doc, line)

        # ── 普通正文 ──
        else:
            add_body(doc, line)

    # 清空未关闭的 table/code
    if in_table and tbl_buf:
        rows = parse_table_rows(tbl_buf)
        add_table(doc, rows)
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    doc.save(out_path)
    print(f'[OK] 已生成: {out_path}')


if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md   = os.path.join(base, '华工科技_策略报告_20260305.md')
    out  = os.path.join(base, '华工科技_策略报告_20260305.docx')
    convert(md, out)
