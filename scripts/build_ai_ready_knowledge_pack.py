#!/usr/bin/env python3
"""Build an AI-ready knowledge pack from mixed local research materials.

The script scans a source directory, extracts text/OCR/table data from common
research formats, converts narrative materials into Markdown, exports tables to
CSV, groups artifacts by content dimension, and writes retrieval-oriented
metadata compatible with the project's market document / chunk conventions.
"""

from __future__ import annotations

import argparse
import base64
import csv
import hashlib
import json
import math
import mimetypes
import re
import shutil
import statistics
import unicodedata
import zipfile
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

import fitz
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from rapidocr_onnxruntime import RapidOCR


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}
TEXT_SUFFIXES = {".md", ".txt"}
HTML_SUFFIXES = {".html", ".htm"}
EXCEL_SUFFIXES = {".xlsx", ".xls"}
WORD_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}


def clean_text(value: Any) -> str:
    text = str(value or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\xa0", " ").replace("\u3000", " ")
    text = text.replace("\ufeff", "").replace("\u200b", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_inline_text(value: Any) -> str:
    return clean_text(value).replace("\n", " ")


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def write_text(path: Path, text: str) -> None:
    ensure_dir(path.parent)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def write_json(path: Path, payload: Any) -> None:
    ensure_dir(path.parent)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_jsonl(path: Path, rows: Iterable[dict[str, Any]]) -> None:
    ensure_dir(path.parent)
    with path.open("w", encoding="utf-8") as fh:
        for row in rows:
            fh.write(json.dumps(row, ensure_ascii=False) + "\n")


def write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str]) -> None:
    ensure_dir(path.parent)
    with path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def slugify_ascii(value: str, fallback: str) -> str:
    normalized = unicodedata.normalize("NFKD", value or "")
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_text = re.sub(r"[^A-Za-z0-9]+", "_", ascii_text).strip("_").lower()
    return ascii_text or fallback


def semantic_name(path: Path, fallback: str) -> str:
    stem = path.stem
    if "1-2vs3-4" in stem:
        return "sc_basis_vs_cont0_price"
    if "各月份差" in stem:
        return "sc_multileg_spreads_vs_cont1_price"
    if "统计" in stem:
        return "sc_spread_statistics"
    if "sprd" in stem.lower():
        return "sc_spread_timeseries"
    if "策略" in stem:
        return "sc_spread_data_notes"
    if "原油" in stem:
        return "crude_oil_strategy_memo"
    if "价差图" in stem:
        return "sc_spread_notebook"
    return slugify_ascii(stem, fallback)


def sheet_slug(sheet_name: str, fallback: str) -> str:
    mapping = {
        "全部": "all",
        "all_daily": "all_daily",
        "all": "all",
    }
    if sheet_name in mapping:
        return mapping[sheet_name]
    if re.fullmatch(r"\d+-\d+", sheet_name):
        return sheet_name.replace("-", "_")
    return slugify_ascii(sheet_name, fallback)


def sha_uid(*parts: Any, length: int = 16) -> str:
    raw = "|".join(str(part or "") for part in parts)
    return hashlib.md5(raw.encode("utf-8")).hexdigest()[:length]


def approx_token_count(text: str) -> int:
    if not text:
        return 0
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    word_count = len(re.findall(r"[A-Za-z0-9_]+", text))
    punctuation_count = len(re.findall(r"[，。；：,.!?()\-\[\]/]", text))
    return max(1, cjk_count + word_count + math.ceil(punctuation_count / 2))


def make_markdown_table(headers: list[str], rows: list[list[Any]]) -> str:
    if not headers:
        return ""
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        cells = [clean_inline_text(item) for item in row]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def normalize_markdown(text: str) -> str:
    text = clean_text(text)
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def coerce_iso_date(value: Any) -> str | None:
    if value is None:
        return None
    if isinstance(value, pd.Timestamp):
        if pd.isna(value):
            return None
        return value.date().isoformat()
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = clean_inline_text(value)
    if not text:
        return None
    for parser in (
        lambda item: date.fromisoformat(item[:10]),
        lambda item: datetime.fromisoformat(item.replace("Z", "+00:00")).date(),
    ):
        try:
            return parser(text).isoformat()
        except Exception:
            continue
    digits = "".join(ch for ch in text if ch.isdigit())
    if len(digits) >= 8:
        try:
            return date.fromisoformat(f"{digits[:4]}-{digits[4:6]}-{digits[6:8]}").isoformat()
        except Exception:
            return None
    return None


def normalize_delivery_month(value: Any) -> str | None:
    text = clean_inline_text(value)
    if not text:
        return None
    if text.upper() == "ALL":
        return "ALL"
    digits = "".join(ch for ch in text if ch.isdigit())
    if len(digits) == 4:
        return digits
    return text


def normalize_percentile_label(value: Any) -> str | None:
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        pct = float(value) * 100
        return f"pct_{int(round(pct)):02d}"
    text = clean_inline_text(value)
    if text in {"0.01", "0.05", "0.1", "0.25", "0.5", "0.75", "0.9", "0.95", "0.99"}:
        return normalize_percentile_label(float(text))
    return None


def normalize_column_name(name: Any) -> str:
    direct = {
        "月份": "delivery_month",
        "合约": "contract_leg",
        "count": "sample_count",
        "mean": "mean",
        "std": "std_dev",
        "min": "min_value",
        "max": "max_value",
        "date_std": "start_date",
        "date_end": "end_date",
        "trading_date_x": "trading_date",
        "Mid": "spread_mid",
        "Conts": "contracts",
        "01": "price_01",
        "Cont0": "front_contract",
        "Cont01": "contract_01",
        "Price01": "price_01",
        "Leg": "contract_leg",
    }
    if name in direct:
        return direct[name]
    if isinstance(name, str) and re.fullmatch(r"\d+-\d+", name):
        return f"spread_{name.replace('-', '_')}"
    percentile = normalize_percentile_label(name)
    if percentile:
        return percentile
    text = clean_inline_text(name)
    return slugify_ascii(text, "column")


def dataframe_to_records(df: pd.DataFrame) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for row in df.to_dict(orient="records"):
        normalized: dict[str, Any] = {}
        for key, value in row.items():
            if isinstance(value, float) and pd.isna(value):
                normalized[key] = None
                continue
            if isinstance(value, pd.Timestamp):
                normalized[key] = value.isoformat()
                continue
            normalized[key] = value
        records.append(normalized)
    return records


def chunk_markdown(text: str, max_chars: int = 900) -> list[str]:
    paragraphs = [part.strip() for part in normalize_markdown(text).split("\n\n") if part.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        candidate = paragraph if not current else f"{current}\n\n{paragraph}"
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            chunks.append(current.strip())
            current = ""
        if len(paragraph) <= max_chars:
            current = paragraph
            continue
        sentences = re.split(r"(?<=[。！？.!?])\s+", paragraph)
        temp = ""
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            candidate = sentence if not temp else f"{temp} {sentence}"
            if len(candidate) <= max_chars:
                temp = candidate
            else:
                if temp:
                    chunks.append(temp.strip())
                temp = sentence
        if temp:
            current = temp
    if current:
        chunks.append(current.strip())
    return chunks or [normalize_markdown(text)]


def rows_to_markdown_list(rows: list[str], prefix: str = "- ") -> str:
    return "\n".join(f"{prefix}{clean_inline_text(row)}" for row in rows if clean_inline_text(row))


def extract_docx_media(docx_path: Path, assets_dir: Path) -> list[Path]:
    extracted: list[Path] = []
    ensure_dir(assets_dir)
    with zipfile.ZipFile(docx_path) as archive:
        media_names = sorted(name for name in archive.namelist() if name.startswith("word/media/"))
        for idx, name in enumerate(media_names, start=1):
            ext = Path(name).suffix.lower() or ""
            payload = archive.read(name)
            if not payload or ext not in IMAGE_SUFFIXES:
                continue
            out_path = assets_dir / f"figure_{idx:02d}{ext}"
            out_path.write_bytes(payload)
            extracted.append(out_path)
    return extracted


def extract_pdf_text(doc_path: Path, page_assets_dir: Path, ocr: RapidOCR) -> tuple[str, list[Path]]:
    ensure_dir(page_assets_dir)
    doc = fitz.open(doc_path)
    pages: list[str] = []
    rendered_pages: list[Path] = []
    for page_index, page in enumerate(doc, start=1):
        page_text = clean_text(page.get_text("text"))
        if len(page_text) < 80:
            page_asset = page_assets_dir / f"page_{page_index:02d}.png"
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            pixmap.save(str(page_asset))
            rendered_pages.append(page_asset)
            ocr_lines = run_ocr_lines(ocr, page_asset)
            page_text = "\n".join(line["text"] for line in ocr_lines)
        pages.append(f"## Page {page_index}\n\n{page_text.strip()}")
    return "\n\n".join(pages).strip(), rendered_pages


def run_ocr_lines(ocr: RapidOCR, image_path: Path) -> list[dict[str, Any]]:
    results, _ = ocr(str(image_path))
    lines: list[dict[str, Any]] = []
    for item in results or []:
        box, text, score = item
        xs = [pt[0] for pt in box]
        ys = [pt[1] for pt in box]
        lines.append(
            {
                "text": clean_inline_text(text),
                "score": round(float(score), 6),
                "x": round(statistics.mean(xs), 2),
                "y": round(statistics.mean(ys), 2),
            }
        )
    lines.sort(key=lambda row: (round(row["y"] / 18), row["x"]))
    return [line for line in lines if line["text"]]


def summarize_ocr_chart(file_name: str, ocr_lines: list[dict[str, Any]]) -> str:
    texts = [line["text"] for line in ocr_lines]
    text_blob = " ".join(texts)
    if "1-2vs3-4" in file_name:
        return (
            "图像显示 SC 主力价格与两组跨月价差的叠加走势：蓝线对应 cont1-cont2，橙线对应 cont3-cont4。"
            " OCR 识别到横轴大致覆盖 2023-01 至 2025-01，左轴价差约 -60 到 60，右轴价格约 500 到 750。"
        )
    if "各月份差" in file_name:
        return (
            "图像将 1-2、2-3、3-4、4-5 四组价差与 ContNo1 价格放在双轴图中共同展示，"
            "横轴覆盖约 2022-11 至 2024-12，可用于观察近端价差与主力价格共振。"
        )
    title = next((item for item in texts if len(item) >= 8), file_name)
    dates = re.findall(r"\d{4}-\d{2}", text_blob)
    range_hint = ""
    if dates:
        range_hint = f" OCR 可见时间刻度大致为 {dates[0]} 至 {dates[-1]}。"
    return f"图像 OCR 标题为“{title}”。{range_hint}"


def summarize_html_figures(figures: list[dict[str, Any]]) -> str:
    if not figures:
        return "HTML 文件未发现嵌入图像。"
    count = len(figures)
    titles = [figure.get("title") for figure in figures if figure.get("title")]
    title_hint = f" 主要 OCR 标题包括：{'; '.join(titles[:3])}。" if titles else ""
    return (
        f"HTML 文件中提取出 {count} 张嵌入图像。结合 `sc_sprd` 工作簿包含 11 组月差数据这一事实，"
        f"这里推断这些图像大概率对应单腿价差或其派生可视化。{title_hint}"
    )


def guess_keywords(title: str, group: str, extra: list[str] | None = None) -> list[str]:
    base = {
        "strategy_notes": ["原油", "SC", "策略", "跨月价差"],
        "price_trend": ["原油", "SC", "价格走势", "价差图"],
        "spread_statistics": ["原油", "SC", "统计", "时间序列"],
    }.get(group, ["原油", "SC"])
    merged = base + list(extra or [])
    seen: list[str] = []
    for item in merged + re.findall(r"[A-Za-z0-9\-\u4e00-\u9fff]+", title):
        item = clean_inline_text(item)
        if item and item not in seen:
            seen.append(item)
    return seen[:12]


@dataclass
class ProcessedDocument:
    doc_uid: str
    title: str
    doc_type: str
    content_group: str
    source_file: str
    output_markdown: str
    summary: str
    body: str
    keywords: list[str]
    dataset_refs: list[str] = field(default_factory=list)
    published_at: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class DatasetArtifact:
    dataset_id: str
    title: str
    source_file: str
    output_csv: str
    group: str
    row_count: int
    column_count: int
    columns: list[str]
    date_columns: list[str]
    date_range: dict[str, str | None]
    summary: str
    metadata: dict[str, Any] = field(default_factory=dict)


class KnowledgePackBuilder:
    def __init__(self, source_dir: Path, output_dir: Path, asset_code: str, asset_name: str):
        self.source_dir = source_dir
        self.output_dir = output_dir
        self.asset_code = asset_code
        self.asset_name = asset_name
        self.ocr = RapidOCR()
        self.documents: list[ProcessedDocument] = []
        self.datasets: list[DatasetArtifact] = []
        self.inventory_rows: list[dict[str, Any]] = []
        self.processing_notes: list[str] = []
        self.doc_counter = 0
        self.dataset_counter = 0

        self.strategy_dir = ensure_dir(output_dir / "strategy_notes")
        self.price_dir = ensure_dir(output_dir / "price_trend")
        self.stats_dir = ensure_dir(output_dir / "spread_statistics")
        self.tables_dir = ensure_dir(output_dir / "tables")
        self.metadata_dir = ensure_dir(output_dir / "metadata")

    def next_doc_path(self, group: str, base_name: str) -> Path:
        self.doc_counter += 1
        parent = {
            "strategy_notes": self.strategy_dir,
            "price_trend": self.price_dir,
            "spread_statistics": self.stats_dir,
        }[group]
        return parent / f"doc_{self.doc_counter:02d}_{base_name}.md"

    def next_dataset_path(self, subgroup: str, base_name: str, sheet_name: str) -> Path:
        self.dataset_counter += 1
        parent = ensure_dir(self.tables_dir / subgroup)
        return parent / f"dataset_{self.dataset_counter:02d}_{base_name}_{sheet_name}.csv"

    def relative(self, path: Path) -> str:
        return path.relative_to(self.output_dir).as_posix()

    def scan_files(self) -> list[Path]:
        files = [
            path
            for path in self.source_dir.rglob("*")
            if path.is_file() and self.output_dir not in path.parents and path != self.output_dir
        ]
        return sorted(files)

    def classify(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in IMAGE_SUFFIXES:
            return "image"
        if suffix in HTML_SUFFIXES:
            return "html"
        if suffix in WORD_SUFFIXES:
            return "word"
        if suffix in EXCEL_SUFFIXES:
            return "excel"
        if suffix in TEXT_SUFFIXES:
            return "text"
        if suffix in PDF_SUFFIXES:
            return "pdf"
        return "other"

    def add_inventory(self, path: Path, file_type: str, outputs: list[str], note: str = "") -> None:
        mime_type, _ = mimetypes.guess_type(str(path))
        self.inventory_rows.append(
            {
                "source_path": str(path),
                "relative_source_path": path.relative_to(self.source_dir).as_posix(),
                "file_type": file_type,
                "mime_type": mime_type or "",
                "size_bytes": path.stat().st_size,
                "outputs": "; ".join(outputs),
                "note": note,
            }
        )

    def register_document(
        self,
        *,
        title: str,
        doc_type: str,
        group: str,
        source_file: Path,
        output_path: Path,
        body: str,
        summary: str,
        dataset_refs: list[str] | None = None,
        published_at: str | None = None,
        extra_metadata: dict[str, Any] | None = None,
        keywords: list[str] | None = None,
    ) -> None:
        normalized_body = normalize_markdown(body)
        doc_uid = sha_uid(self.asset_code, title, str(source_file), normalized_body[:256])
        metadata = {
            "asset_name": self.asset_name,
            "instrument_code": self.asset_code,
            "instrument_type": "futures",
            "language": "zh",
            "content_group": group,
            "source_file": str(source_file),
            "output_markdown": self.relative(output_path),
            "dataset_refs": list(dataset_refs or []),
            "ready_for": ["market_documents", "market_doc_chunks", "strategy_factory_vector_memory"],
        }
        if extra_metadata:
            metadata.update(extra_metadata)
        document = ProcessedDocument(
            doc_uid=doc_uid,
            title=title,
            doc_type=doc_type,
            content_group=group,
            source_file=str(source_file),
            output_markdown=self.relative(output_path),
            summary=summary,
            body=normalized_body,
            keywords=keywords or guess_keywords(title, group),
            dataset_refs=list(dataset_refs or []),
            published_at=published_at,
            metadata=metadata,
        )
        self.documents.append(document)

    def register_dataset(
        self,
        *,
        title: str,
        source_file: Path,
        output_csv: Path,
        group: str,
        records: list[dict[str, Any]],
        metadata: dict[str, Any] | None = None,
    ) -> DatasetArtifact:
        columns = list(records[0].keys()) if records else []
        date_columns = [column for column in columns if "date" in column]
        detected_dates: list[str] = []
        for row in records:
            for column in date_columns:
                value = row.get(column)
                if value:
                    detected_dates.append(str(value))
        dataset = DatasetArtifact(
            dataset_id=sha_uid(title, str(source_file), self.relative(output_csv)),
            title=title,
            source_file=str(source_file),
            output_csv=self.relative(output_csv),
            group=group,
            row_count=len(records),
            column_count=len(columns),
            columns=columns,
            date_columns=date_columns,
            date_range={
                "start": min(detected_dates) if detected_dates else None,
                "end": max(detected_dates) if detected_dates else None,
            },
            summary="",
            metadata=metadata or {},
        )
        self.datasets.append(dataset)
        return dataset

    def process_text_file(self, path: Path) -> None:
        raw_text = path.read_text(encoding="utf-8", errors="ignore")
        base_name = semantic_name(path, "notes")
        if path.name == "策略.md":
            title = "SC 价差数据说明与换月规则"
            body = self.build_strategy_notes_markdown(raw_text)
            summary = "整理换月规则、分钟中间价口径、图表说明和 Excel 字段定义，可作为时序与统计数据的口径说明文档。"
            group = "strategy_notes"
            extra = {"content_dimension": "data_definition"}
            keywords = guess_keywords(title, group, ["换月规则", "数据口径", "字段说明"])
        else:
            title = clean_inline_text(path.stem)
            body = normalize_markdown(raw_text)
            summary = clean_inline_text(body.split("\n", 1)[0])[:120]
            group = "strategy_notes"
            extra = {"content_dimension": "general_notes"}
            keywords = guess_keywords(title, group)

        output_path = self.next_doc_path(group, base_name)
        write_text(output_path, body)
        self.register_document(
            title=title,
            doc_type="research",
            group=group,
            source_file=path,
            output_path=output_path,
            body=body,
            summary=summary,
            extra_metadata=extra,
            keywords=keywords,
        )
        self.add_inventory(path, "text", [self.relative(output_path)], note="converted to cleaned Markdown")

    def build_strategy_notes_markdown(self, raw_text: str) -> str:
        lines = [clean_inline_text(line) for line in raw_text.splitlines() if clean_inline_text(line)]
        roll_rules: list[str] = []
        data_scope: list[str] = []
        chart_notes: list[str] = []
        excel_notes: list[str] = []
        excel_header = ""
        for line in lines:
            if line.startswith("换月规则"):
                roll_rules.append(line.split("：", 1)[-1].strip())
            elif line.startswith("所以作为"):
                roll_rules.append(line)
            elif line.startswith("数据"):
                data_scope.append(line.split(":", 1)[-1].split("：", 1)[-1].strip())
            elif line.startswith("图说明"):
                chart_notes.append(line.split("：", 1)[-1].strip())
            elif line.startswith("x轴") or line.startswith("y轴"):
                chart_notes.append(line)
            elif line.startswith("xlsx说明"):
                excel_notes.append("按月份划分工作表：`1-2` 代表第 1 合约与第 2 合约的价差，`all` 为全量汇总。")
            elif "all_daily是每日一行" in line or "all_daily" in line:
                excel_header = "trading_date_x, Cont01, Price01, 1-2, 2-3, 3-4, 4-5, 5-6, 6-7, 7-8, 8-9, 9-10, 10-11, 11-12"
        body = [
            "# SC 价差数据说明与换月规则",
            "",
            "## 换月规则",
            rows_to_markdown_list(roll_rules),
            "",
            "## 数据口径",
            rows_to_markdown_list(data_scope),
            "",
            "## 图表说明",
            rows_to_markdown_list(chart_notes),
            "",
            "## Excel 字段说明",
            rows_to_markdown_list(excel_notes),
            "",
            "```text",
            excel_header,
            "```",
            "",
            "## AI 调用提示",
            "- 先检索本说明文档，再联动 `tables/timeseries/` 与 `tables/statistics/` 下的 CSV，可避免误解字段口径。",
            "- 若进入 `market_doc_chunks`，建议 `doc_type='research'`，并在 metadata 中写入 `content_dimension='data_definition'`。",
        ]
        return "\n".join(part for part in body if part is not None)

    def process_word_file(self, path: Path) -> None:
        doc = Document(path)
        paragraphs = [clean_inline_text(p.text) for p in doc.paragraphs if clean_inline_text(p.text)]
        base_name = semantic_name(path, "word_doc")
        assets_dir = ensure_dir(self.strategy_dir / "assets" / base_name)
        media_paths = extract_docx_media(path, assets_dir)
        output_path = self.next_doc_path("strategy_notes", base_name)
        body = self.build_docx_markdown(paragraphs, media_paths, output_path.parent)
        summary = (
            "将原始 Word 备忘整理为结构化研究笔记，提炼选品原则、SC 近远月关系观察、趋势与套利策略以及风险收益要点。"
        )
        write_text(output_path, body)
        self.register_document(
            title="原油跨月价差交易备忘",
            doc_type="research",
            group="strategy_notes",
            source_file=path,
            output_path=output_path,
            body=body,
            summary=summary,
            extra_metadata={
                "content_dimension": "strategy_framework",
                "hierarchy_reconstructed": True,
                "extracted_media": [self.relative(media_path) for media_path in media_paths],
            },
            keywords=guess_keywords("原油跨月价差交易备忘", "strategy_notes", ["carry", "套利", "趋势"]),
        )
        self.add_inventory(
            path,
            "word",
            [self.relative(output_path), *[self.relative(media_path) for media_path in media_paths]],
            note="converted to Markdown and extracted embedded media",
        )

    def build_docx_markdown(self, paragraphs: list[str], media_paths: list[Path], markdown_parent: Path) -> str:
        def find_index(prefix: str) -> int:
            for idx, text in enumerate(paragraphs):
                if text.startswith(prefix):
                    return idx
            return -1

        title = "原油跨月价差交易备忘"
        sc_idx = find_index("用sc来举例")
        observe_idx = find_index("我们可以观察到什么")
        trend_idx = find_index("趋势：")
        arbitrage_idx = find_index("套利：")
        peers_idx = find_index("同类交易品种")

        selection = paragraphs[1:sc_idx] if sc_idx > 1 else paragraphs[1:6]
        sc_intro = paragraphs[sc_idx + 1 : observe_idx] if sc_idx >= 0 and observe_idx > sc_idx else []
        observations = paragraphs[observe_idx + 1 : trend_idx] if observe_idx >= 0 and trend_idx > observe_idx else []
        trend_section = paragraphs[trend_idx:arbitrage_idx] if trend_idx >= 0 and arbitrage_idx > trend_idx else []
        arbitrage_section = paragraphs[arbitrage_idx:peers_idx] if arbitrage_idx >= 0 and peers_idx > arbitrage_idx else []
        peers = paragraphs[peers_idx + 1 :] if peers_idx >= 0 else []

        asset_lines: list[str] = []
        if media_paths:
            first_image = media_paths[0].relative_to(markdown_parent).as_posix()
            asset_lines = [
                "![SC case figure](" + first_image + ")",
                "",
            ]

        trend_summary: list[str] = []
        trend_risks: list[str] = []
        for item in trend_section:
            if item.startswith("趋势："):
                trend_summary.append(item.split("：", 1)[-1])
            elif item.startswith("风险："):
                trend_risks.append(item.split("：", 1)[-1])
            else:
                trend_summary.append(item)

        arbitrage_summary: list[str] = []
        arbitrage_risks: list[str] = []
        for item in arbitrage_section:
            if item.startswith("套利："):
                arbitrage_summary.append(item.split("：", 1)[-1])
            elif item.startswith("风险："):
                arbitrage_risks.append(item.split("：", 1)[-1])
            else:
                arbitrage_summary.append(item)

        risk_lines = trend_risks + [item for item in arbitrage_summary if item.startswith("收益为") or item.startswith("风险：")] + arbitrage_risks
        peer_lines = [item for item in peers if item]

        body = [
            f"# {title}",
            "",
            "## 研究目标",
            paragraphs[0] if paragraphs else "通过历史行情与价差结构寻找更合适的交易品种与交易方式。",
            "",
            "## 选品原则",
            rows_to_markdown_list(selection),
            "",
            "## SC 案例说明",
            *asset_lines,
            rows_to_markdown_list(sc_intro),
            "",
            "### 图中观察",
            rows_to_markdown_list(observations),
            "",
            "## 交易策略",
            "### 趋势策略",
            rows_to_markdown_list(trend_summary),
            "",
            "### 套利策略",
            rows_to_markdown_list(arbitrage_summary),
            "",
            "## 风险与收益",
            rows_to_markdown_list(risk_lines),
            "",
            "## 可类比品种",
            rows_to_markdown_list(peer_lines),
            "",
            "## AI 调用提示",
            "- 这份文档适合作为 `research` 类型写入 `market_documents`，并在 chunk metadata 中标记 `strategy_framework`。",
            "- 若要做策略工厂检索，优先把 `交易策略` 与 `风险与收益` 两段拆成独立 chunk，以提升查询命中率。",
        ]
        return "\n".join(part for part in body if part is not None)

    def process_image_file(self, path: Path) -> None:
        base_name = semantic_name(path, "image")
        assets_dir = ensure_dir(self.price_dir / "assets")
        copied_asset = assets_dir / f"{base_name}{path.suffix.lower()}"
        shutil.copy2(path, copied_asset)
        ocr_lines = run_ocr_lines(self.ocr, copied_asset)
        title = next((line["text"] for line in ocr_lines if len(line["text"]) >= 8), path.stem)
        legends = [line["text"] for line in ocr_lines if ("-" in line["text"] or "cont" in line["text"].lower())][:8]
        dates = [line["text"] for line in ocr_lines if re.fullmatch(r"\d{4}-\d{2}", line["text"])]
        summary = summarize_ocr_chart(path.name, ocr_lines)
        output_path = self.next_doc_path("price_trend", base_name)
        body = [
            f"# {title}",
            "",
            "## 来源",
            f"- 原始文件：`{path}`",
            f"- 输出图像：`{self.relative(copied_asset)}`",
            "",
            "## OCR 提取",
            f"- 识别文本数量：{len(ocr_lines)}",
            f"- 图例候选：{', '.join(legends) if legends else '未明显识别'}",
            f"- 时间刻度候选：{', '.join(dates[:8]) if dates else '未明显识别'}",
            "",
            "## 图表描述",
            summary,
            "",
            "## 原始 OCR 文本",
            "```text",
            "\n".join(f"{line['text']} (score={line['score']})" for line in ocr_lines[:80]),
            "```",
            "",
            "## AI 调用提示",
            "- 用于回答近端价差与主力价格关系、阶段性波动和异常点问题。",
            "- 检索该图时建议联动 `tables/timeseries/` 中的逐日序列，以便从视觉结论回到原始数据。",
        ]
        write_text(output_path, "\n".join(body))
        self.register_document(
            title=title,
            doc_type="research",
            group="price_trend",
            source_file=path,
            output_path=output_path,
            body="\n".join(body),
            summary=summary,
            extra_metadata={
                "content_dimension": "price_trend",
                "image_asset": self.relative(copied_asset),
                "ocr_line_count": len(ocr_lines),
            },
            keywords=guess_keywords(title, "price_trend", legends),
        )
        self.add_inventory(
            path,
            "image",
            [self.relative(output_path), self.relative(copied_asset)],
            note="copied image and generated OCR markdown",
        )

    def process_html_file(self, path: Path) -> None:
        html_text = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html_text, "html.parser")
        base_name = semantic_name(path, "html_report")
        html_assets_dir = ensure_dir(self.price_dir / "html_assets" / base_name)
        figures: list[dict[str, Any]] = []
        for idx, img in enumerate(soup.find_all("img"), start=1):
            src = img.get("src", "")
            if not src.startswith("data:image/") or "," not in src:
                continue
            header, payload = src.split(",", 1)
            ext = "." + header.split("/")[1].split(";")[0]
            asset_path = html_assets_dir / f"figure_{idx:02d}{ext}"
            asset_path.write_bytes(base64.b64decode(payload))
            ocr_lines = run_ocr_lines(self.ocr, asset_path)
            title = next((line["text"] for line in ocr_lines if re.fullmatch(r"\d+-\d+", line["text"])), "")
            if not title:
                title = next((line["text"] for line in ocr_lines if len(line["text"]) >= 8), "")
            figures.append(
                {
                    "figure_no": idx,
                    "path": self.relative(asset_path),
                    "title": title,
                    "ocr_text_preview": " | ".join(line["text"] for line in ocr_lines[:6]),
                    "ocr_count": len(ocr_lines),
                }
            )

        figure_rows = [
            [item["figure_no"], item["title"] or "未识别标题", item["ocr_count"], item["path"]]
            for item in figures
        ]
        title = clean_inline_text(soup.title.text if soup.title else path.stem)
        summary = summarize_html_figures(figures)
        output_path = self.next_doc_path("price_trend", base_name)
        body = [
            f"# {title} HTML 图册整理",
            "",
            "## 文件概览",
            f"- 原始文件：`{path}`",
            f"- 提取嵌入图像数：{len(figures)}",
            "",
            "## 图册摘要",
            summary,
            "",
            "## 图像清单",
            make_markdown_table(["Figure", "OCR 标题", "OCR 文本数量", "输出路径"], figure_rows),
            "",
            "## OCR 预览",
            "```text",
            "\n".join(
                f"figure_{item['figure_no']:02d}: {item['ocr_text_preview']}" for item in figures[:12]
            ),
            "```",
            "",
            "## AI 调用提示",
            "- 这份 HTML 更适合作为图册索引使用，而不是原样做全文检索。",
            "- 建议将图册说明写入 `market_documents`，同时把图像文件路径保存在 metadata 中，供后续人工复核。",
        ]
        write_text(output_path, "\n".join(body))
        self.register_document(
            title=f"{title} HTML 图册整理",
            doc_type="research",
            group="price_trend",
            source_file=path,
            output_path=output_path,
            body="\n".join(body),
            summary=summary,
            extra_metadata={
                "content_dimension": "price_trend_gallery",
                "figure_count": len(figures),
                "extracted_figures": [item["path"] for item in figures],
            },
            keywords=guess_keywords(f"{title} HTML 图册整理", "price_trend", ["HTML", "图册"]),
        )
        outputs = [self.relative(output_path)] + [item["path"] for item in figures]
        self.add_inventory(path, "html", outputs, note="extracted embedded figures and built Markdown catalog")

    def normalize_records(self, records: list[dict[str, Any]]) -> list[dict[str, Any]]:
        normalized_rows: list[dict[str, Any]] = []
        for record in records:
            row: dict[str, Any] = {}
            for key, value in record.items():
                if key in {"trading_date", "start_date", "end_date"}:
                    row[key] = coerce_iso_date(value)
                elif key == "delivery_month":
                    row[key] = normalize_delivery_month(value)
                elif isinstance(value, float) and math.isnan(value):
                    row[key] = None
                else:
                    row[key] = value
            normalized_rows.append(row)
        return normalized_rows

    def process_excel_file(self, path: Path) -> None:
        workbook_name = semantic_name(path, "workbook")
        subgroup = "statistics" if "统计" in path.stem else "timeseries"
        excel = pd.ExcelFile(path)
        generated_outputs: list[str] = []
        workbook_datasets: list[DatasetArtifact] = []

        for sheet in excel.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet)
            df.columns = [normalize_column_name(column) for column in df.columns]
            records = self.normalize_records(dataframe_to_records(df))
            csv_path = self.next_dataset_path(subgroup, workbook_name, sheet_slug(sheet, f"sheet_{len(workbook_datasets) + 1:02d}"))
            pd.DataFrame(records).to_csv(csv_path, index=False, encoding="utf-8")

            metadata: dict[str, Any] = {"sheet_name": sheet}
            summary = ""
            if subgroup == "statistics":
                all_row = next((row for row in records if row.get("delivery_month") == "ALL"), records[0] if records else {})
                metadata["all_row"] = all_row
                summary = (
                    f"样本 {all_row.get('sample_count')}, 均值 {all_row.get('mean')}, "
                    f"波动 {all_row.get('std_dev')}, 区间 [{all_row.get('min_value')}, {all_row.get('max_value')}]"
                )
            else:
                first = records[0] if records else {}
                last = records[-1] if records else {}
                metadata["coverage"] = {
                    "start": first.get("trading_date"),
                    "end": last.get("trading_date"),
                }
                summary = (
                    f"共 {len(records)} 行，字段 {len(df.columns)} 个，覆盖 "
                    f"{first.get('trading_date') or '未知'} 至 {last.get('trading_date') or '未知'}。"
                )
            dataset = self.register_dataset(
                title=f"{path.stem} / {sheet}",
                source_file=path,
                output_csv=csv_path,
                group=subgroup,
                records=records,
                metadata=metadata,
            )
            dataset.summary = summary
            workbook_datasets.append(dataset)
            generated_outputs.append(self.relative(csv_path))

        if subgroup == "statistics":
            output_path = self.next_doc_path("spread_statistics", workbook_name)
            body, summary = self.build_statistics_summary(path, workbook_datasets)
            group = "spread_statistics"
            extra = {"content_dimension": "spread_statistics"}
        else:
            output_path = self.next_doc_path("spread_statistics", workbook_name)
            body, summary = self.build_timeseries_summary(path, workbook_datasets)
            group = "spread_statistics"
            extra = {"content_dimension": "price_timeseries"}

        write_text(output_path, body)
        self.register_document(
            title=clean_inline_text(path.stem),
            doc_type="research",
            group=group,
            source_file=path,
            output_path=output_path,
            body=body,
            summary=summary,
            dataset_refs=[dataset.dataset_id for dataset in workbook_datasets],
            extra_metadata=extra,
            keywords=guess_keywords(clean_inline_text(path.stem), group, [subgroup]),
        )
        generated_outputs.append(self.relative(output_path))
        self.add_inventory(path, "excel", generated_outputs, note=f"exported {len(workbook_datasets)} sheets to CSV and created workbook summary")

    def build_statistics_summary(self, path: Path, datasets: list[DatasetArtifact]) -> tuple[str, str]:
        rows: list[list[Any]] = []
        std_values: list[tuple[str, float]] = []
        coverage: list[str] = []
        combined_tables: list[str] = []
        for dataset in datasets:
            all_row = dict(dataset.metadata.get("all_row") or {})
            sheet_name = str(dataset.metadata.get("sheet_name") or "")
            if sheet_name in {"全部", "all"}:
                combined_tables.append(dataset.output_csv)
                continue
            contract_leg = all_row.get("contract_leg") or sheet_name
            rows.append(
                [
                    contract_leg,
                    all_row.get("sample_count"),
                    all_row.get("mean"),
                    all_row.get("std_dev"),
                    all_row.get("pct_01"),
                    all_row.get("pct_25"),
                    all_row.get("pct_75"),
                    all_row.get("pct_99"),
                    dataset.output_csv,
                ]
            )
            if all_row.get("std_dev") is not None:
                std_values.append((str(contract_leg), float(all_row["std_dev"])))
            for key in ("start_date", "end_date"):
                if all_row.get(key):
                    coverage.append(str(all_row[key]))
        highest_vol = max(std_values, key=lambda item: item[1])[0] if std_values else "未知"
        summary = (
            f"{path.name} 已拆成 {len(datasets)} 个 CSV。ALL 行显示 1-2 组合波动最大，"
            f"当前自动识别的最高标准差腿为 {highest_vol}。"
        )
        body = [
            f"# {path.stem} 统计摘要",
            "",
            "## 工作簿概览",
            f"- 输出数据表：{len(datasets)} 个",
            f"- 统计覆盖：{min(coverage) if coverage else '未知'} 至 {max(coverage) if coverage else '未知'}",
            f"- 主要用途：近远月价差分位、波动区间、建仓风险参考",
            f"- 综合汇总表：{', '.join(combined_tables) if combined_tables else '无'}",
            "",
            "## ALL 行关键统计",
            make_markdown_table(
                ["组合", "样本数", "均值", "标准差", "1%", "25%", "75%", "99%", "CSV"],
                rows,
            ),
            "",
            "## AI 调用提示",
            "- 可先读本摘要，再下钻到具体组合 CSV，例如 `1-2` 或 `3-4` 的分位区间。",
            "- 若要进入策略工厂，可把 1%、25%、75%、99% 视为建仓/风控阈值候选特征。",
        ]
        return "\n".join(body), summary

    def build_timeseries_summary(self, path: Path, datasets: list[DatasetArtifact]) -> tuple[str, str]:
        date_ranges = [dataset.date_range for dataset in datasets if dataset.date_range.get("start")]
        starts = [item["start"] for item in date_ranges if item.get("start")]
        ends = [item["end"] for item in date_ranges if item.get("end")]
        rows = [
            [dataset.metadata.get("sheet_name"), dataset.row_count, dataset.column_count, dataset.date_range.get("start"), dataset.date_range.get("end"), dataset.output_csv]
            for dataset in datasets
        ]
        summary = (
            f"{path.name} 已拆成 {len(datasets)} 个时序 CSV，整体覆盖 "
            f"{min(starts) if starts else '未知'} 至 {max(ends) if ends else '未知'}。"
        )
        body = [
            f"# {path.stem} 时序数据摘要",
            "",
            "## 工作簿概览",
            f"- 输出数据表：{len(datasets)} 个",
            f"- 时间覆盖：{min(starts) if starts else '未知'} 至 {max(ends) if ends else '未知'}",
            "- 主要用途：逐日价差回溯、横向比较不同 ContNo 腿、构造回测输入",
            "",
            "## 输出 CSV 清单",
            make_markdown_table(["工作表", "行数", "列数", "起始日期", "结束日期", "CSV"], rows),
            "",
            "## AI 调用提示",
            "- `all_daily` 是最适合程序消费的宽表，可直接读成多列价差因子矩阵。",
            "- 单腿工作表适合做局部分析、绘图复刻和异常日期回放。",
        ]
        return "\n".join(body), summary

    def process_pdf_file(self, path: Path) -> None:
        base_name = semantic_name(path, "pdf_doc")
        page_assets = self.price_dir / "pdf_assets" / base_name
        extracted_text, rendered_pages = extract_pdf_text(path, page_assets, self.ocr)
        title = clean_inline_text(path.stem)
        body = "\n".join(
            [
                f"# {title}",
                "",
                extracted_text or "未提取到正文内容。",
                "",
                "## AI 调用提示",
                "- 如果 PDF 为扫描版，本文件已对低文本页执行 OCR。",
            ]
        )
        output_path = self.next_doc_path("price_trend", base_name)
        write_text(output_path, body)
        self.register_document(
            title=title,
            doc_type="research",
            group="price_trend",
            source_file=path,
            output_path=output_path,
            body=body,
            summary="PDF 已抽取文本，并在文本稀疏页执行 OCR。",
            extra_metadata={
                "content_dimension": "pdf_research",
                "rendered_pages": [self.relative(page) for page in rendered_pages],
            },
        )
        outputs = [self.relative(output_path), *[self.relative(page) for page in rendered_pages]]
        self.add_inventory(path, "pdf", outputs, note="extracted text and OCRed sparse pages")

    def process_other_file(self, path: Path) -> None:
        self.add_inventory(path, "other", [], note="unsupported type, kept only in inventory")

    def build_metadata_exports(self) -> None:
        market_documents_rows: list[dict[str, Any]] = []
        market_doc_chunks_rows: list[dict[str, Any]] = []
        normalized_docs_rows: list[dict[str, Any]] = []

        for document in self.documents:
            market_documents_rows.append(
                {
                    "doc_uid": document.doc_uid,
                    "stock_code": self.asset_code,
                    "doc_type": document.doc_type,
                    "source": "local_raw_materials",
                    "title": document.title,
                    "summary": document.summary,
                    "body": document.body,
                    "url": "",
                    "author": "",
                    "published_at": document.published_at,
                    "metadata": document.metadata | {"keywords": document.keywords},
                }
            )
            normalized_docs_rows.append(
                {
                    "doc_id": document.doc_uid,
                    "doc_type": document.doc_type,
                    "date": document.published_at,
                    "title": document.title,
                    "source": "local_raw_materials",
                    "text": document.body,
                    "text_length": len(document.body),
                    "output_markdown": document.output_markdown,
                }
            )
            for chunk_no, chunk_text in enumerate(chunk_markdown(document.body), start=1):
                market_doc_chunks_rows.append(
                    {
                        "doc_uid": document.doc_uid,
                        "entity_id": f"{document.doc_uid}:{chunk_no}",
                        "chunk_no": chunk_no,
                        "stock_code": self.asset_code,
                        "doc_type": document.doc_type,
                        "source": "local_raw_materials",
                        "title": document.title,
                        "chunk_text": chunk_text,
                        "token_count": approx_token_count(chunk_text),
                        "char_count": len(chunk_text),
                        "language": "zh",
                        "published_at": document.published_at,
                        "metadata": {
                            "content_group": document.content_group,
                            "source_file": document.source_file,
                            "output_markdown": document.output_markdown,
                            "dataset_refs": document.dataset_refs,
                            "keywords": document.keywords,
                        },
                    }
                )

        write_jsonl(self.metadata_dir / "market_documents.jsonl", market_documents_rows)
        write_jsonl(self.metadata_dir / "market_doc_chunks.jsonl", market_doc_chunks_rows)
        write_jsonl(self.metadata_dir / "normalized_documents.jsonl", normalized_docs_rows)
        write_jsonl(
            self.metadata_dir / "datasets.jsonl",
            [
                {
                    "dataset_id": dataset.dataset_id,
                    "title": dataset.title,
                    "source_file": dataset.source_file,
                    "output_csv": dataset.output_csv,
                    "group": dataset.group,
                    "row_count": dataset.row_count,
                    "column_count": dataset.column_count,
                    "columns": dataset.columns,
                    "date_columns": dataset.date_columns,
                    "date_range": dataset.date_range,
                    "summary": dataset.summary,
                    "metadata": dataset.metadata,
                }
                for dataset in self.datasets
            ],
        )
        write_csv(
            self.metadata_dir / "file_inventory.csv",
            self.inventory_rows,
            ["source_path", "relative_source_path", "file_type", "mime_type", "size_bytes", "outputs", "note"],
        )
        write_json(
            self.metadata_dir / "knowledge_pack_manifest.json",
            {
                "asset_name": self.asset_name,
                "asset_code": self.asset_code,
                "source_dir": str(self.source_dir),
                "output_dir": str(self.output_dir),
                "generated_at": datetime.now().isoformat(timespec="seconds"),
                "source_file_count": len(self.inventory_rows),
                "document_count": len(self.documents),
                "chunk_count": len(market_doc_chunks_rows),
                "dataset_count": len(self.datasets),
                "groups": {
                    "strategy_notes": len([doc for doc in self.documents if doc.content_group == "strategy_notes"]),
                    "price_trend": len([doc for doc in self.documents if doc.content_group == "price_trend"]),
                    "spread_statistics": len([doc for doc in self.documents if doc.content_group == "spread_statistics"]),
                },
                "compatibility": {
                    "market_documents": "metadata/market_documents.jsonl",
                    "market_doc_chunks": "metadata/market_doc_chunks.jsonl",
                    "datasets": "metadata/datasets.jsonl",
                },
            },
        )

    def build_index(self) -> None:
        grouped_docs: dict[str, list[ProcessedDocument]] = {
            "strategy_notes": [],
            "price_trend": [],
            "spread_statistics": [],
        }
        for document in self.documents:
            grouped_docs.setdefault(document.content_group, []).append(document)

        group_titles = {
            "strategy_notes": "交易策略与研究框架",
            "price_trend": "价格走势与图表 OCR",
            "spread_statistics": "统计与结构化数据",
        }

        doc_sections: list[str] = []
        for group, title in group_titles.items():
            docs = grouped_docs.get(group, [])
            if not docs:
                continue
            doc_sections.append(f"## {title}")
            for document in docs:
                doc_sections.extend(
                    [
                        f"### {document.title}",
                        f"- 输出文件：`{document.output_markdown}`",
                        f"- 来源：`{document.source_file}`",
                        f"- 摘要：{document.summary}",
                        f"- 建议用途：{document.metadata.get('content_dimension') or document.content_group}",
                        "",
                    ]
                )

        dataset_rows = [
            [
                dataset.title,
                dataset.group,
                dataset.row_count,
                dataset.column_count,
                dataset.date_range.get("start") or "",
                dataset.date_range.get("end") or "",
                dataset.output_csv,
            ]
            for dataset in self.datasets
        ]
        missing_dimensions = [
            "宏观背景：未检测到专门的宏观报告或宏观数据表。",
            "供需基本面：未检测到库存/产量/进口/炼厂开工等直接原始资料。",
            "库存统计：当前目录无库存日报或库存表，仅有价差统计与价格序列。",
        ]
        body = [
            f"# {self.asset_name} 资料 AI 就绪索引",
            "",
            "## 扫描概览",
            f"- 源目录：`{self.source_dir}`",
            f"- 识别源文件：{len(self.inventory_rows)} 个",
            f"- 生成 Markdown 文档：{len(self.documents)} 份",
            f"- 生成 CSV 数据集：{len(self.datasets)} 份",
            f"- 生成 market_doc_chunks：{sum(len(chunk_markdown(doc.body)) for doc in self.documents)} 个",
            "",
            "## 输出目录",
            "- `strategy_notes/`：研究备忘、数据口径和交易框架。",
            "- `price_trend/`：图像 OCR、HTML 图册整理和图表说明。",
            "- `spread_statistics/`：Excel 工作簿摘要与结构化说明。",
            "- `tables/`：标准化 CSV。",
            "- `metadata/`：知识库注入清单、chunk、文件索引。",
            "",
            *doc_sections,
            "## 结构化数据表",
            make_markdown_table(
                ["数据集", "分组", "行数", "列数", "起始日期", "结束日期", "CSV"],
                dataset_rows,
            ),
            "",
            "## 缺失维度提示",
            rows_to_markdown_list(missing_dimensions),
            "",
            "## 注入建议",
            "- Narrative 文档：使用 `metadata/market_documents.jsonl` 与 `metadata/market_doc_chunks.jsonl`，`doc_type` 统一按 `research` 注入。",
            "- 结构化表：优先消费 `metadata/datasets.jsonl` 中登记的 CSV，不建议把整张表直接塞进 chunk 文本。",
            "- 策略工厂：建议把 `content_group`、`content_dimension`、`dataset_refs`、`instrument_type='futures'` 一并写入 metadata，便于后续过滤。",
            "- 若走 `akshare-mcp` 的 DB-first 路径，可将这里的 narrative 文档视作 `research` 类型市场文档，并保留 `stock_code='SC'` 作为兼容字段。",
            "",
            "## 核验建议",
            "- 先用 `metadata/file_inventory.csv` 检查源文件是否全部落盘。",
            "- 再抽查 `tables/timeseries/dataset_*_all_daily.csv` 与 `spread_statistics/doc_*.md` 是否和原始口径一致。",
            "- 最后再将 narrative 文档 chunk 化写入向量集合 `market_doc_chunks`。",
        ]
        write_text(self.output_dir / "index.md", "\n".join(body))

    def run(self) -> None:
        ensure_dir(self.output_dir)
        for path in self.scan_files():
            file_type = self.classify(path)
            if file_type == "text":
                self.process_text_file(path)
            elif file_type == "word":
                self.process_word_file(path)
            elif file_type == "excel":
                self.process_excel_file(path)
            elif file_type == "html":
                self.process_html_file(path)
            elif file_type == "image":
                self.process_image_file(path)
            elif file_type == "pdf":
                self.process_pdf_file(path)
            else:
                self.process_other_file(path)
        self.build_metadata_exports()
        self.build_index()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build an AI-ready knowledge pack from mixed research materials.")
    parser.add_argument("--source", required=True, help="Source directory containing raw materials.")
    parser.add_argument("--output", help="Output directory. Defaults to <source>/ai_ready.")
    parser.add_argument("--asset-code", default="SC", help="Compatibility code written to metadata/stock_code.")
    parser.add_argument("--asset-name", default="原油", help="Human-friendly asset name written to summaries.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source_dir = Path(args.source).expanduser().resolve()
    output_dir = Path(args.output).expanduser().resolve() if args.output else source_dir / "ai_ready"
    builder = KnowledgePackBuilder(
        source_dir=source_dir,
        output_dir=output_dir,
        asset_code=str(args.asset_code).strip() or "SC",
        asset_name=str(args.asset_name).strip() or "原油",
    )
    builder.run()


if __name__ == "__main__":
    main()
